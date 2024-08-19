from flask import Flask, request, jsonify
import json
from spire.doc import Document, FileFormat
from docx import Document as DocxDocument
import docx2pdf
import os
import uuid

app = Flask(__name__)

def process_modified_file(input_docx, output_dir, base_name):
    output_docx = os.path.join(output_dir, f'{base_name}.docx')
    output_pdf = os.path.join(output_dir, f'{base_name}.pdf')
    
    document = DocxDocument(input_docx)
    target_string = "Evaluation Warning: The document was created with Spire.Doc for Python."
    for paragraph in document.paragraphs:
        if target_string in paragraph.text:
            paragraph.text = paragraph.text.replace(target_string, "")
    
    document.save(output_docx)
    docx2pdf.convert(output_docx, output_pdf)
    
    os.remove(output_docx)
    os.remove(input_docx)
    
    return base_name  # Return the base name used in the generated filenames

@app.route('/replace-text', methods=['POST'])
def replace_text_in_docx():
    try:
        # Recibir los datos JSON y el UUID desde la solicitud
        json_data = request.get_json()
        uuid = request.args.get('uuid')  # Obtener el UUID del parámetro de consulta

        # Crear el directorio basado en el UUID si no existe
        output_dir = os.path.join(os.getcwd(), uuid)
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

        # Recorrer cada documento y sus reemplazos en el JSON
        for doc_path, replacements in json_data.items():
            # Crear la ruta del archivo modificado en el directorio del UUID
            base_name = os.path.basename(doc_path).replace('.docx', '')
            modified_doc_path = os.path.join(output_dir, "Reemplazado_" + os.path.basename(doc_path))

            # Cargar el documento
            document = Document()
            document.LoadFromFile(doc_path)
            
            # Reemplazar texto en el documento
            for search_text, replace_text in replacements.items():
                document.Replace(search_text, replace_text, False, False)
            
            # Guardar el documento con los cambios
            document.SaveToFile(modified_doc_path, FileFormat.Docx2016)
            document.Close()
            
            # Procesar el archivo modificado y obtener el nombre base
            base_name = process_modified_file(modified_doc_path, output_dir, base_name)

        return jsonify({"message": "Text replacement and processing completed successfully.", "unique_id": base_name}), 200

    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/valida-campos', methods=['POST'])
def valida_campos():
    data = request.json
    file_path = data.get('file_path')
    campos = data.get('campos')

    if not file_path or not campos:
        return jsonify({'error': 'file_path y campos son requeridos'}), 400

    try:
        # Crear un objeto Document
        document = Document()
        
        # Cargar el documento Word docx o doc
        document.LoadFromFile(file_path)

        # Obtener el texto del documento
        texto_documento = document.GetText()

        # Buscar campos faltantes
        campos_faltantes = [campo for campo in campos if campo not in texto_documento]

        if campos_faltantes:
            return jsonify({'campos_faltantes': campos_faltantes})
        else:
            return jsonify({'message': 'Todos los campos están presentes.'})

    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000)
